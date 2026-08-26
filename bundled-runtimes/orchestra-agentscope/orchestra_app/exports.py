from __future__ import annotations

import html
from io import BytesIO
from typing import Any

from .models import RunSnapshot


def _sections(
    snapshot: RunSnapshot,
    artifacts: list[dict[str, Any]],
    evidence: list[dict[str, Any]],
) -> list[tuple[str, str]]:
    agent_reports = []
    for artifact in artifacts:
        if artifact["kind"].endswith("_report"):
            agent_reports.append(f"{artifact['title']}\n{artifact['content']}")
    evidence_lines = []
    for item in evidence:
        source = item["source_name"]
        interface = item.get("interface_name") or item["tool_name"]
        observed = item.get("observed_at") or "未标注"
        url = item.get("source_url") or ""
        evidence_lines.append(
            f"{source} | {interface} | 数据日期 {observed} | 抓取 {item['retrieved_at']}"
            + (f" | {url}" if url else ""),
        )
    return [
        ("运行信息", f"议题：{snapshot.topic}\n版本：v{snapshot.revision}\n模式：{snapshot.mode}\n状态：{snapshot.status}"),
        ("研究计划", snapshot.plan or "未生成"),
        ("阶段成果", "\n\n".join(agent_reports) or "未生成"),
        ("共识与分歧", snapshot.consensus or "未生成"),
        ("正式投决报告", snapshot.decision or "未生成"),
        ("证据链", "\n".join(evidence_lines) or "本次运行没有已记录的外部数据证据。"),
    ]


def build_docx(
    snapshot: RunSnapshot,
    artifacts: list[dict[str, Any]],
    evidence: list[dict[str, Any]],
) -> bytes:
    from docx import Document
    from docx.shared import Pt

    document = Document()
    document.add_heading("Orchestra 投资决策报告", level=0)
    document.add_paragraph(f"Run ID: {snapshot.id}")
    for title, content in _sections(snapshot, artifacts, evidence):
        document.add_heading(title, level=1)
        for line in content.splitlines() or [""]:
            paragraph = document.add_paragraph(line)
            paragraph.style.font.size = Pt(10.5)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_pdf(
    snapshot: RunSnapshot,
    artifacts: list[dict[str, Any]],
    evidence: list[dict[str, Any]],
) -> bytes:
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    buffer = BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title=f"Orchestra 投决报告 {snapshot.id}",
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ChineseTitle",
        parent=styles["Title"],
        fontName="STSong-Light",
        fontSize=20,
        leading=28,
        alignment=TA_CENTER,
        spaceAfter=12,
    )
    heading_style = ParagraphStyle(
        "ChineseHeading",
        parent=styles["Heading1"],
        fontName="STSong-Light",
        fontSize=13,
        leading=20,
        spaceBefore=10,
        spaceAfter=6,
    )
    body_style = ParagraphStyle(
        "ChineseBody",
        parent=styles["BodyText"],
        fontName="STSong-Light",
        fontSize=9.5,
        leading=16,
        spaceAfter=4,
    )
    story = [
        Paragraph("Orchestra 投资决策报告", title_style),
        Paragraph(f"Run ID: {html.escape(snapshot.id)}", body_style),
        Spacer(1, 6),
    ]
    for index, (title, content) in enumerate(_sections(snapshot, artifacts, evidence)):
        if index == 2:
            story.append(PageBreak())
        story.append(Paragraph(html.escape(title), heading_style))
        for line in content.splitlines() or [""]:
            story.append(Paragraph(html.escape(line) or "&nbsp;", body_style))
    document.build(story)
    return buffer.getvalue()
